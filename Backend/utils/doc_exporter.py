from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document


# PDF formating
def export_pdf(project_name: str, docs: str, path: str):
    pdf = SimpleDocTemplate(path, pagesize=A4)

    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(project_name, styles["Title"]))
    story.append(Spacer(1, 20))

    for line in docs.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 12))

    pdf.build(story)


# Document formating
def export_docx(project_name: str, docs: str, path: str):
    document = Document()

    document.add_heading(project_name, level=1)

    for line in docs.split("\n"):
        if line.strip():
            document.add_paragraph(line)

    document.save(path)
